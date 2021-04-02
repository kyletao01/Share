import docx
from docx.shared import Inches
from docx_func import insert_fig


def gen_doc(top,inst_top):
    Template_doc=docx.Document("Format-Template.docx")
    doc=Template_doc

    #=============================================================
    #Chapter 
    #AIPU system registers
    #-------------------------------------------------------------
    paragragh = doc.add_paragraph('\nAIPU system registers',style='Heading 2')
    paragragh.runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE) # change page
    #-------------------------------------------------------------

    #=============================================================
    #AIPU architectural system register summary
    #-------------------------------------------------------------
    doc.add_paragraph('AIPU architectural system register summary',style='Heading 3')
    doc.add_paragraph('This section identifies the architectural system registers implemented in the AIPU core.',style='BodyText')

    # Add table
    doc.add_paragraph("Table x-x "+inst_top.full_name,style='FigCaption')
    #doc.add_paragraph("Table {STYLEREF 2\\s }-{SEQ Table \\* ARABIC \\s 2 } "+inst_top.full_name+" Registers Group",style='FigCaption')
    titles    = ['Address','Name','Reset value','Description']
    table = doc.add_table(rows=1, cols=len(titles))
    # Table format
    table.style = 'Table Grid'
    # set Table Title
    title_cells = table.rows[0].cells
    for i in range(len(titles)):
        cell=title_cells[i]
        cell.text = titles[i]
        for p in cell.paragraphs:
            p.style = 'ARM table head'
    # fill table content
    for block in inst_top.blocks:
        for reg_name in inst_top.blocks[block].regs:
            reg=inst_top.blocks[block].regs[reg_name]
            add = reg.offset 
            name = reg.full_name.replace(" Register", "")
            #name = reg.full_name.rstrip(' Register')
            reset = reg.reset_value or "0x0"
            des = reg.description or reg.full_name
            data=[add,name,reset,des]
            row_cells = table.add_row().cells
            for i in range(len(titles)):
                cell =row_cells[i] 
                cell.text = data[i]
                for p in cell.paragraphs:
                    p.style = 'ARM table text'
    # set columns width
    col_width = [0.87, 1.94, 1.22, 2.94]
    for  i,col  in  enumerate(table.columns):
        for  cell in col.cells:
            cell.width=Inches(col_width[i])
    doc.add_paragraph('\n')
    doc.paragraphs[-1].runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)
    #=============================================================
                

    #=============================================================
    #Registers by functional group
    #-------------------------------------------------------------


    #=============================================================
    #Chapter 
    #AIPU Registers Group
    #-------------------------------------------------------------
    paragragh= doc.add_paragraph('\nControl Registers Group',style='Heading 2')
    paragragh.runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE) # change page
    #-------------------------------------------------------------

    def add_reg(reg):
        doc.add_paragraph(reg.full_name,style='Heading 3')

        # Add fig
        doc.add_paragraph("The following figure shows the {} bit assignments".format(reg.full_name),style='BodyText')

        field_list = list()
        lsb_list = list()
        for field_name in reg.fields:
            field=reg.fields[field_name]
            # if(':' in str(field.bits)):
            #     lsb = int(field.bits.split(":")[1]) 
            # else:
            #     lsb = int(field.bits) 
            field_list.append(field.name)
            lsb_list.append(field.lsb)
        insert_fig(doc,field_list,lsb_list)
        doc.add_paragraph("Figure x-x "+reg.full_name+" bit assignments",style='FigCaption')
        # doc.add_paragraph('\n')

        # Add table
        doc.add_paragraph("The following table shows the {} bit assignments".format(reg.full_name),style='BodyText')
        doc.add_paragraph("Table x-x "+reg.full_name+" bit assignments",style='FigCaption')
        titles    = ['Bits','Name','Default value','Access','Description']
        table = doc.add_table(rows=1, cols=len(titles))
        # Table format
        table.style = 'Table Grid'
        # set Table Title
        title_cells = table.rows[0].cells
        for i in range(len(titles)):
            cell=title_cells[i]
            cell.text = titles[i]
            for p in cell.paragraphs:
                p.style = 'ARM table head'
        # fill table content
        for field_name in reg.fields:
            field=reg.fields[field_name]
            bits = field.bits 
            name = field.full_name
            reset = field.reset_value or "-"
            access = field.access
            des = field.description or "-"
            data=[bits,name,reset,access,des]
            for i,val in enumerate(data):
                data[i]=str(val)
            row_cells = table.add_row().cells
            for i in range(len(titles)):
                cell =row_cells[i] 
                cell.text = data[i]
                for p in cell.paragraphs:
                    p.style = 'ARM table text'
        doc.add_paragraph('\n')
        # set columns width
        col_width = [0.53, 1.22, 1.06, 0.81,3.37]
        for  i,col  in  enumerate(table.columns):
            for  cell in col.cells:
                cell.width=Inches(col_width[i])

        if reg.description != "":
            doc.add_paragraph(reg.description,style='BodyText')

        doc.paragraphs[-1].runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)

        #=============================================================

    for block in top.blocks:
        for reg_name in top.blocks[block].regs:
            reg=top.blocks[block].regs[reg_name]
            add_reg(reg)

    return doc

if __name__ == '__main__':

    import os
    os.chdir(os.path.dirname(os.path.realpath(__file__)))
    from extract_info import *

    top,inst_top = get_demo_reg_tree()
    doc=gen_doc(top,inst_top)
    doc.save('Register.docx')#保存文档