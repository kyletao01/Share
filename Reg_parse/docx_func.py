import os
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.xmlchemy import BaseOxmlElement, ZeroOrOne, ZeroOrMore, OxmlElement
from docx.oxml.ns import qn

def seq_paragraph(paragraph): 
    '''
    set a auto-order sequence for table  or figure
    Usage:
    paragraph = doc.add_paragraph('Figure ', style='Caption') 
    seq_paragraph(paragraph)
    '''
    run = run = paragraph.add_run() 
    r = run._r 
    fldChar = OxmlElement('w:fldChar') 
    fldChar.set(qn('w:fldCharType'), 'begin') 
    r.append(fldChar) 
    instrText = OxmlElement('w:instrText') 
    instrText.text = ' SEQ Table \\* ARABIC' 
    r.append(instrText) 
    fldChar = OxmlElement('w:fldChar') 
    fldChar.set(qn('w:fldCharType'), 'end') 
    r.append(fldChar)

def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def insert_fig(doc,field_list,lsb_list):
    '''
    field_list=['ALL_BYP','Reserved','RELU_BYP','MUL_PSRC_DTYPE','MUL_PSRC_PREC',]
    lsb_list=[31,12,5,1,0]
    '''
    table = doc.add_table(rows=2, cols=32)
    table.style = 'Normal Table'

    title_cells = table.rows[0].cells
    for i in range(32):
        cell=title_cells[i]
        cell.text = str(31-i)
        for p in cell.paragraphs:
            p.style = 'ARM page number R'

    start  = 31
    for index,i in enumerate(lsb_list):
        end = i
        if start != end:
            table.cell(0,31-start+1).merge(table.cell(0,31-end-1))
            cell = table.cell(0,31-start+1) 
            cell.text=str()
            for p in cell.paragraphs:
                p.style = 'ARM page number R'
        table.cell(1,31-start).merge(table.cell(1,31-end))
        cell = table.cell(1,31-start) 
        cell.text=field_list[index]
        for p in cell.paragraphs:
            p.style = 'FigCaption'
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_border(
                cell,
                top={"val": "single"},
                bottom={"val": "single"},
                left={"val": "single" },
                right={"val": "single"},
            )
        start = i-1

if __name__ == '__main__':
    os.chdir(os.path.dirname(os.path.realpath(__file__)))
    #output_file = os.path.join(os.getcwd(),"output.xml")

    Template_doc=docx.Document("Format-Template.docx")
    # doc=docx.Document()#创建一个Document对象
    doc=Template_doc

    paragraph = doc.add_paragraph('Figure ', style='Caption') 
    Table(paragraph)
